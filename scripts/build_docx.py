#!/usr/bin/env python3
"""
Build the MoC Daily Cultural Digest .docx from its canonical markdown, using
python-docx. The canonical markdown is the single source of truth; this
script is a faithful renderer, not an editor.

FORMATTING VALUES BELOW ARE NOT GUESSES: they were reverse-engineered
directly from the OOXML of real delivered editions by unzipping each .docx
and inspecting word/document.xml, word/styles.xml, and word/theme/theme1.xml.
If a future reference edition contradicts these values, re-derive from that
edition the same way, not from this comment.

REVISION HISTORY:
- 2026-07-19/20: first pass, built against the 19 July real edition. Used
  bilingual "English (Arabic)" commission labels, Aptos 12pt, hyperlinks
  #0563C1, explicit point-based paragraph spacing.
- 2026-07-21: corrected against the 16 July real edition, which is
  chronologically EARLIER than the 19 July one but was confirmed by the
  user as the actual target to match, with two explicit clarifications:
  (a) commission labels should be plain English, no Arabic -- the whole
  document gets translated into Arabic as a separate downstream step, so
  the English edition's own labels don't need to carry Arabic; (b) base
  font size is 11pt, not 12pt; (c) paragraph spacing should use Word's
  built-in "No Spacing" style (0pt before/after) with real blank paragraphs
  between sections for visual separation, not explicit point values.

Confirmed current values (2026-07-21, font corrected same day):
- Body font: Times New Roman, 11pt. The reference edition's docDefaults
  explicitly set Times New Roman as the base font; most individual runs
  additionally carried a `w:rFonts asciiTheme="majorBidi"` override which
  resolves to "Aptos Display" via the document's theme -- an earlier pass
  on 2026-07-21 assumed that per-run theme override was the one that
  actually renders, and built against Aptos. Directly corrected by the
  user: the real, intended font is Times New Roman throughout, not Aptos.
- Section headings ("Headlines, <date>", "Saudi Arabia/Regional", "Global"):
  bold, color #4A86E8 (blue).
- "Negative Articles" heading: bold, color #EE0000 (red) -- CONFIRMED
  DIFFERENT from the other two section headings in the 16 July reference;
  this was not visible in the 19 July reference and is a real, newly
  confirmed per-section color convention, not a guess.
- Commission subheadings (e.g. "Heritage:"): bold, black, plain English,
  colon-terminated -- no Arabic, no parentheses.
- "Risks" subsection heading: bold, color #C00000 (dark red).
- "Opportunities" subsection heading: bold, color #A66500 (dark gold).
- Numbered item headlines inside Risks/Opportunities: bold, black.
- Hyperlinks: color #0000FF (plain blue), single underline -- CORRECTED
  from #0563C1, which was confirmed accurate for the 19 July edition but
  not the 16 July one; #0000FF is what the 16 July reference actually uses.
- Paragraph spacing: Word's built-in "No Spacing" paragraph style (0pt
  before/after) throughout, INCLUDING bullets -- not explicit point values.
  Visual separation between sections comes from real blank paragraphs
  ("line drops"), not from paragraph-level spacing settings.

WHY python-docx INSTEAD OF DIRECT XML ON A PINNED TEMPLATE:
The conflict-monitoring skill this repo was adapted from edits a pre-branded,
SHA-256-pinned template directly via XML, for exact brand fidelity. This repo
does not yet have a real branded MoC Word template to pin, so it builds a
clean, professionally formatted document from scratch instead, styled to
match the real reference editions' actual formatting values above. Once a
real branded template (letterhead, logo, house styles) is supplied, swap
this script for the pinned-template direct-XML-edit approach the same way
the conflict-monitoring skill does -- see templates/README.md.

CANONICAL MARKDOWN FORMAT this script expects (see SKILL.md Stage 3):

    # Headlines, <DATE>

    ## Saudi Arabia/Regional
    - Headline one
    - Headline two

    ## Negative Articles
    - Headline one

    ## Global
    - Headline one

    # Saudi Arabia/Regional

    ## General:
    - Free-form analytical bullet ending in a plain citation. ([Outlet](url))

    ## Heritage:
    - ...

    # Negative Articles
    - Bullets go directly here -- NO commission subheadings in this section.
    - ([Outlet](url))

    # Global

    ## Museums:
    - ...

    # Risks and Opportunities

    ## Risks

    1. **Short bold headline naming the specific risk**
    Analytical paragraph.
    Source: [Outlet](url), [Outlet](url)
    Consideration: paragraph.

    2. **A second risk, if supported**
    Paragraph.
    Source: [Outlet](url)
    Consideration: paragraph.

    ## Opportunities

    1. **Short bold headline naming the specific opportunity**
    Paragraph.
    Source: [Outlet](url), [Outlet](url), [Outlet](url)
    Consideration: paragraph.

Note: each of Risks/Opportunities restarts its own numbering at 1, and each
subsection can have one or more items -- the count isn't fixed.

Usage:
    python3 scripts/build_docx.py reports/MoC_Digest_2026-07-19.md \\
        --output reports/MoC_Digest_2026-07-19.docx
"""

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT_NAME = "Times New Roman"
FONT_SIZE = 11

# Colors, confirmed from the 16 July real reference edition's OOXML (see
# module docstring above for the revision history and rationale).
COLOR_HEADING_BLUE = RGBColor(0x4A, 0x86, 0xE8)
COLOR_NEGATIVE_RED = RGBColor(0xEE, 0x00, 0x00)
COLOR_BLACK = RGBColor(0x00, 0x00, 0x00)
COLOR_RISK_RED = RGBColor(0xC0, 0x00, 0x00)
COLOR_OPPORTUNITY_GOLD = RGBColor(0xA6, 0x65, 0x00)
COLOR_HYPERLINK = "0000FF"  # hex string form, used in raw OxmlElement below

# Per-section heading color -- confirmed 2026-07-21 that Negative Articles
# gets its own distinct red, not the same blue as the other two sections.
SECTION_HEADING_COLORS = {
    "Saudi Arabia/Regional": COLOR_HEADING_BLUE,
    "Negative Articles": COLOR_NEGATIVE_RED,
    "Global": COLOR_HEADING_BLUE,
}

MD_LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
NUMBERED_ITEM_RE = re.compile(r"^\s*(\d+)\.\s*(.*)$")
BOLD_HEADLINE_RE = re.compile(r"^\*\*(.+?)\*\*\s*$")

REQUIRED_SECTIONS = ["Saudi Arabia/Regional", "Negative Articles", "Global"]
# Sections that use commission subheadings in the full-summary block.
# Negative Articles is deliberately excluded, per confirmed real production.
SECTIONS_WITH_SUBHEADINGS = {"Saudi Arabia/Regional", "Global"}


def add_hyperlink(paragraph, text, url):
    """
    Insert a real clickable hyperlink run into a python-docx paragraph.
    python-docx has no native hyperlink API, so this builds the relationship
    and run XML directly -- the standard recipe for this library. Color and
    underline match the 16 July reference edition's actual hyperlink runs
    (#0000FF, single underline) exactly.
    """
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), COLOR_HYPERLINK)
    rpr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)

    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), FONT_NAME)
    rFonts.set(qn("w:hAnsi"), FONT_NAME)
    rpr.append(rFonts)

    run.append(rpr)
    text_el = OxmlElement("w:t")
    text_el.set(qn("xml:space"), "preserve")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)

    paragraph._p.append(hyperlink)
    return hyperlink


def add_mixed_text_with_links(paragraph, text_with_md_links, bold=False):
    """
    Adds text to a paragraph that may contain one or more markdown-style
    [label](url) links (bullet citations, or comma-separated Source lines).
    Plain text is added as normal runs; each link becomes a real hyperlink.
    """
    remainder = text_with_md_links
    while True:
        match = MD_LINK_RE.search(remainder)
        if not match:
            if remainder:
                run = paragraph.add_run(remainder)
                run.bold = bold
            break
        if remainder[: match.start()]:
            run = paragraph.add_run(remainder[: match.start()])
            run.bold = bold
        add_hyperlink(paragraph, match.group(1), match.group(2))
        remainder = remainder[match.end() :]


def set_base_style(doc):
    """Sets the Normal style to Times New Roman 11pt, matching the 16 July reference edition."""
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(FONT_SIZE)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rFonts.set(qn("w:ascii"), FONT_NAME)
    rFonts.set(qn("w:hAnsi"), FONT_NAME)


def no_spacing_paragraph(doc):
    """
    Every non-bulleted paragraph uses Word's built-in "No Spacing" style
    (0pt before/after), confirmed against the 16 July reference. Visual
    separation between sections comes from real blank paragraphs, not from
    paragraph-level spacing settings -- see add_blank_line() below.
    """
    return doc.add_paragraph(style="No Spacing")


def add_blank_line(doc):
    """
    A real empty paragraph, used to create the visual "line drop" between
    sections that the 16 July reference achieves this way, rather than
    through paragraph spacing settings.
    """
    doc.add_paragraph(style="No Spacing")


def add_heading_paragraph(doc, text, section_name_for_color):
    """
    Section/title-level heading: bold, colored per SECTION_HEADING_COLORS
    (Negative Articles is red, others are blue), No Spacing style.
    """
    p = no_spacing_paragraph(doc)
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = SECTION_HEADING_COLORS.get(section_name_for_color, COLOR_HEADING_BLUE)
    return p


def add_subheading_paragraph(doc, text, color=COLOR_BLACK):
    """Commission subheading: bold, black by default, No Spacing style, plain English."""
    p = no_spacing_paragraph(doc)
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    return p


def add_bullet_paragraph(doc, text_with_links):
    """Body bullet: List Bullet style, No Spacing-equivalent (no explicit override), real hyperlinks embedded."""
    p = doc.add_paragraph(style="List Bullet")
    add_mixed_text_with_links(p, text_with_links)
    return p


def split_h1_blocks(md_text):
    """Splits markdown into (h1_name, h1_body_lines) blocks, in order."""
    blocks = []
    current_name = None
    current_lines = []
    for line in md_text.split("\n"):
        if line.startswith("# ") and not line.startswith("## "):
            if current_name is not None:
                blocks.append((current_name, current_lines))
            current_name = line[2:].strip()
            current_lines = []
        else:
            current_lines.append(line)
    if current_name is not None:
        blocks.append((current_name, current_lines))
    return blocks


def parse_headline_block(h1_lines):
    """Parses the 'Headlines, <date>' block into {section: [headline, ...]}."""
    result = {}
    current_h2 = None
    for line in h1_lines:
        if line.startswith("## "):
            current_h2 = line[3:].strip()
            result.setdefault(current_h2, [])
        elif line.strip().startswith("- ") and current_h2:
            result[current_h2].append(line.strip()[2:].strip())
    return result


def parse_full_summary_block(h1_name, h1_lines):
    """
    Parses a full-summary section into {commission_label: [bullet, ...]}.
    Sections in SECTIONS_WITH_SUBHEADINGS use ## commission subheadings
    (plain English, e.g. "Heritage:"); Negative Articles (and any other
    section not in that set) has bullets directly under the H1, with no
    subheading -- stored under a single implicit "" key.
    """
    result = {}
    if h1_name in SECTIONS_WITH_SUBHEADINGS:
        current_h2 = None
        for line in h1_lines:
            if line.startswith("## "):
                current_h2 = line[3:].strip()
                result.setdefault(current_h2, [])
            elif line.strip().startswith("- ") and current_h2:
                result[current_h2].append(line.strip()[2:].strip())
    else:
        result[""] = []
        for line in h1_lines:
            if line.strip().startswith("- "):
                result[""].append(line.strip()[2:].strip())
    return result


def parse_risks_opportunities_block(h1_lines):
    """
    Parses the Risks and Opportunities block into:
    {
        "risks": [ {"headline": str, "paragraph": str, "source": str, "consideration": str}, ... ],
        "opportunities": [ ... ],
    }
    Each ## subsection contains one or more numbered items. An item starts
    at a line matching NUMBERED_ITEM_RE and runs until the next numbered
    item or the end of the subsection.
    """
    result = {"risks": [], "opportunities": []}
    current_h2 = None
    item_lines = []

    def flush_item():
        if not item_lines:
            return
        text = "\n".join(item_lines).strip()
        if not text:
            return

        lines = text.split("\n")
        first_line = lines[0]
        num_match = NUMBERED_ITEM_RE.match(first_line)
        headline_raw = num_match.group(2).strip() if num_match else first_line.strip()
        bold_match = BOLD_HEADLINE_RE.match(headline_raw)
        headline = bold_match.group(1).strip() if bold_match else headline_raw

        body = "\n".join(lines[1:]).strip()
        source_match = re.search(r"Source:\s*(.+?)(?:\nConsideration:|$)", body, re.DOTALL)
        consideration_match = re.search(r"Consideration:\s*(.+)", body, re.DOTALL)

        if source_match:
            paragraph = body[: source_match.start()].strip()
        else:
            paragraph = body.strip()

        item = {
            "headline": headline,
            "paragraph": paragraph,
            "source": source_match.group(1).strip() if source_match else "",
            "consideration": consideration_match.group(1).strip() if consideration_match else "",
        }
        key = current_h2.lower() if current_h2 else "risks"
        result.setdefault(key, []).append(item)

    for line in h1_lines:
        if line.startswith("## "):
            flush_item()
            item_lines = []
            current_h2 = line[3:].strip()
        elif NUMBERED_ITEM_RE.match(line.strip()) and item_lines:
            flush_item()
            item_lines = [line.strip()]
        elif line.strip():
            item_lines.append(line.strip())
    flush_item()

    return result


def parse_markdown(md_text):
    """
    Parses the full canonical markdown into:
    {
        "title": "Headlines, 19 July 2026",
        "headline_bullets": {section: [headline, ...]},
        "sections": {section: {commission_label: [bullet, ...]}},
        "risks_and_opportunities": {"risks": [...], "opportunities": [...]},
    }
    """
    blocks = split_h1_blocks(md_text)
    result = {
        "title": "",
        "headline_bullets": {},
        "sections": {},
        "risks_and_opportunities": {},
    }
    seen_headline_block = False

    for h1_name, h1_lines in blocks:
        if h1_name.startswith("Headlines,"):
            result["title"] = h1_name
            result["headline_bullets"] = parse_headline_block(h1_lines)
            seen_headline_block = True
        elif h1_name in REQUIRED_SECTIONS:
            result["sections"][h1_name] = parse_full_summary_block(h1_name, h1_lines)
        elif h1_name == "Risks and Opportunities":
            result["risks_and_opportunities"] = parse_risks_opportunities_block(h1_lines)

    if not seen_headline_block:
        print("WARNING: no 'Headlines, <date>' block found in markdown", file=sys.stderr)

    return result


def build_docx(parsed, output_path):
    doc = Document()
    set_base_style(doc)

    # Title -- same styling as the section headings (blue, bold)
    add_heading_paragraph(doc, parsed["title"] or "MoC Daily Cultural Digest", "Saudi Arabia/Regional")
    add_blank_line(doc)

    # Headline bullets FIRST (confirmed real ordering), three sections
    for section_name in REQUIRED_SECTIONS:
        bullets = parsed["headline_bullets"].get(section_name, [])
        add_heading_paragraph(doc, section_name, section_name)
        for bullet_text in bullets:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(bullet_text)
        add_blank_line(doc)

    doc.add_page_break()

    # Full summaries, grouped by section then (where applicable) commission
    for section_name in REQUIRED_SECTIONS:
        add_heading_paragraph(doc, section_name, section_name)

        commissions = parsed["sections"].get(section_name, {})
        for commission_name, bullets in commissions.items():
            if not bullets:
                continue
            if commission_name:  # empty string = no subheading (Negative Articles)
                add_subheading_paragraph(doc, commission_name)

            for bullet_text in bullets:
                add_bullet_paragraph(doc, bullet_text)

        add_blank_line(doc)

    doc.add_page_break()

    # Risks and Opportunities: multiple numbered items per subsection
    add_heading_paragraph(doc, "Risks and Opportunities", "Saudi Arabia/Regional")
    add_blank_line(doc)

    ro = parsed.get("risks_and_opportunities", {})
    subsection_colors = {"risks": COLOR_RISK_RED, "opportunities": COLOR_OPPORTUNITY_GOLD}

    for label in ["risks", "opportunities"]:
        items = ro.get(label, [])
        if not items:
            continue

        add_subheading_paragraph(doc, label.capitalize(), color=subsection_colors[label])

        for idx, item in enumerate(items, start=1):
            headline_p = no_spacing_paragraph(doc)
            run = headline_p.add_run(f"{idx}. {item['headline']}")
            run.bold = True
            run.font.color.rgb = COLOR_BLACK

            if item.get("paragraph"):
                p = no_spacing_paragraph(doc)
                p.add_run(item["paragraph"])

            if item.get("source"):
                p = no_spacing_paragraph(doc)
                run = p.add_run("Source: ")
                run.bold = True
                add_mixed_text_with_links(p, item["source"])

            if item.get("consideration"):
                p = no_spacing_paragraph(doc)
                run = p.add_run("Consideration: ")
                run.bold = True
                p.add_run(item["consideration"])

            add_blank_line(doc)

    doc.save(output_path)
    print(f"Wrote {output_path}")


def main():
    parser = argparse.ArgumentParser(description="Build the MoC Daily Cultural Digest .docx")
    parser.add_argument("markdown_path", help="Path to the canonical markdown digest")
    parser.add_argument("--output", required=True, help="Output .docx path")
    args = parser.parse_args()

    md_text = Path(args.markdown_path).read_text(encoding="utf-8")
    parsed = parse_markdown(md_text)
    build_docx(parsed, args.output)


if __name__ == "__main__":
    main()
